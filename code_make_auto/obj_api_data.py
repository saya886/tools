import requests
import json
import sys
import random
import re
import docx
import os
from docx.shared import Pt

def make_docx(page_header_name, file_name, data_list):
        # 生成docx
        document  = docx.Document('template.docx')
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        # 格式化页眉
        # arg_txt = page_header_name
        # A4 纸的页眉中文最大长度
        # max_len = 70
        # res_txt = arg_txt + " "*(70-(len(arg_txt)*2)) 

        paragraph.text = page_header_name
        random.shuffle(data_list)
        for i in data_list:
            document.add_paragraph(i,style=None)

        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(10)
        document.save(file_name)   
        print("make file " + file_name + " success..........")

class get_str:

    def __init__(self):
        pass

    def search_repositories(self, q,language):
        # 搜索请求
        r = requests.get('https://api.github.com/search/repositories?q={}+language:{}&page=1&per_page=10'.format(q,language), headers=self.headers)
        try:
            res = json.loads(r.text)
        except:
            return {}
        # 随机请求
        index_page = random.randint(0,9)
        r = requests.get('https://api.github.com/search/repositories?q={}+language:{}&page={}&per_page=100'.format(q,language,index_page), headers=self.headers)
        print("request search_repositories url is " + 'https://api.github.com/search/repositories?q={}+language:{}&page={}&per_page=100&sort=stars&order=desc'.format(q,language,index_page))
        try:
            res = json.loads(r.text)
        except:
            res = {}
        return res

    def get_json(self, url):
        # 返回json
        r = requests.get(url, headers=self.headers)
        try:
            res = json.loads(r.text)
        except:
            res = []
        return res 

    def get_string(self, url):
        # 请求url 返回去除空格的字符串 及其行数
        
        get_str = requests.get(url, headers=self.headers_2).text
        str_arg = ""

        # 正则表达式匹配 敏感信息
        get_str = re.sub(r'[a-zA-z]+://[^\s]*', "", get_str)
        get_str = re.sub(r'(C|c)opyright', "", get_str)
        get_str = re.sub(r'\(c\)', "", get_str)
        get_str = re.sub(r'\(C\)', "", get_str)
        get_str = re.sub(r'All rights reserved', "", get_str)
        get_str = re.sub(r'all rights reserved', "", get_str)
        get_str = re.sub(r'\/\*(\s|.)*?\*\/', "", get_str)
        get_str = re.sub(r'[\u4e00-\u9fa5]+', "", get_str)
        # get_str = re.sub(r'(\s\s)', "", get_str)
        for i in get_str.splitlines():
            if i != "":
                str_arg  = str_arg + i + "\n"
        line_count = len(str_arg.splitlines())
        return str_arg,line_count

    def universal_tree_filter_keywords(self, in_url):
        # 输入仓库url 广度优先遍历仓库 根据关键字过滤后缀
        self.deepfirst_search(in_url)
        return self.res_list

    def deepfirst_search(self, root):
        # 深度优先请求
        res_json = self.get_json(root)
        for j in res_json:
            # 累计的行数计数器 大于 需要的行数
            print(j["path"])
            print("当前的行数为 " + str(self.gl_line_count))
            if self.gl_line_count > self.need_gl_line_count:
                print("抓取给定行数完成 行数为 " + str(self.gl_line_count))
                break
            if j["type"] == "dir":
                self.deepfirst_search(j["url"])
            if j["type"] == "file":
                if j["path"].find(self.filter_keyword) != -1:
                    print("************** 完成一个文件 读取 " + j["path"])
                    str_arg,line_count = self.get_string(j["url"])

                    # 增加行数计数器
                    if ((self.gl_line_count + line_count) - self.need_gl_line_count) >= 200:
                        print(" 最后一个文件太大了 再次请求")
                        continue
                   
                    print("!!!!!!!!!!!!!!!!!!!!!!!当前的行数为 " + str(self.gl_line_count))
                    self.res_list.append(str_arg)
                    
                    # # 每添加一次 都写入一次文件  并且 比较一下文件的大小
                    # t_name = "t.docx"
                    # make_docx("headerheaderheaderheader",str(self.count)+t_name,self.res_list)

                    # # 如果写入一个文件后  文件的大小不在 规定的范围内  弹出最后一个元素
                    # if os.path.getsize(str(self.count)+t_name) > 41415:
                    #     print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!warning final file is not stisifai" + str(self.count))
                    #     self.count = self.count +1
                    #     self.res_list.pop()
                    # else:
                    #     self.gl_line_count = self.gl_line_count + line_count
                    
                    self.gl_line_count = self.gl_line_count + line_count

                    

    def return_str_line(self,filter_keyword,search_keyword,lan,need_gl_line_count):
        # 当前的行数
        self.gl_line_count = 0
        # 需要的行数
        self.need_gl_line_count = need_gl_line_count
        self.headers = {'Authorization': 'token 4497921e3e96e86c299b5147ad746a4553c539b7'}
        self.headers_2 = {'Authorization': 'token 4497921e3e96e86c299b5147ad746a4553c539b7', 'accept': 'application/vnd.github.VERSION.raw'}
        # 为返回的字符串 列表
        self.res_list = []
        self.filter_keyword = filter_keyword
        self.search_keyword = search_keyword
        self.lan = lan
        self.count = 0

        res = self.search_repositories(self.search_keyword, self.lan)
        print("当次搜索的条目数 " + str(len(res["items"])))
        # 存储数据的列表
        data_str_list = []
        for i in res["items"]:
            print("开始遍历 的仓库 名称" + i["full_name"])
            print("开始遍历 的仓库 URL" + i["url"])
            print("当前的行数为 " + str(self.gl_line_count))
            print("需要的行数为 " + str(self.need_gl_line_count))
            if self.gl_line_count >= self.need_gl_line_count:
                print("抓取给定行数完成 行数为 " + str(self.gl_line_count))
                break
            ress_list = self.universal_tree_filter_keywords(i["url"]+"/contents")
            print("完成遍历 " + i["full_name"] + i["url"])
            data_str_list = data_str_list + ress_list
            
        return data_str_list