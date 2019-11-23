#coding:utf-8
import docx
import re
# from data import str_arg

def clear_blank(str_arg):
    res = ""
    for i in str_arg.splitlines():
        if i != "":
            res  = res + i + "\n"
    return res




def save_docx(page_header_name,file_name,data_list):

    document  = docx.Document('template.docx')
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    # 格式化页眉
    arg_txt = page_header_name
    max_len = 70
    res_txt = arg_txt + " "*(70-(len(arg_txt)*2))

    paragraph.text = res_txt
    for i in data_list:
        document.add_paragraph(i,style=None)
    document.save(file_name)

save_docx("123","123",[1,23])