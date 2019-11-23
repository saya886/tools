import time
start = time.time()
import os
import csv
import docx
import re
import sys
import random
from docx_red import save_docx
from obj_api_data import get_str
from docx.shared import Pt
import pickle


# def is_continue(second_sount):

class get_doc():
    def __init__(self, NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN):
        self.NEED  = 10
        self.FILTER_KEYWORD = FILTER_KEYWORD
        self.SEARCH_KEYWORD = SEARCH_KEYWORD
        self.LAN = LAN

        self.need_line_count = need_line_count
    
        # 读取包含数据列表的csv文件
        self.f3 = open("data.csv", encoding='utf-8')
        self.file_date_list = csv.reader(self.f3,delimiter='$')

    def get_str(self):
        # 根据行数  请求对应行数的数据
        get_str_instance = get_str()
        return get_str_instance.return_str_line(self.FILTER_KEYWORD,self.SEARCH_KEYWORD,self.LAN,self.need_line_count)
    
    def read_csv(self):
        # 根据行数  请求对应行数的数据
        # 读取包含数据列表的csv文件
        with open("data.csv", encoding='utf-8') as f:
            data_list = csv.reader(f,delimiter=',')
            for row in data_list:
                file_name_1 = row[0] + " " + row[2] + " " + row[3] + " 说明书.docx"
                print(file_name_1)


    def make_docx(self, page_header_name, file_name, data_list):
        # 生成docx
        document  = docx.Document('template.docx')
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        # 格式化页眉
        # arg_txt = page_header_name
        # A4 纸的页眉中文最大长度
        # max_len = 70
        # res_txt = arg_txt + " "*(70-(len(arg_txt)*2)) 

        paragraph.text = page_header_name
        random.shuffle(data_list)
        for i in data_list:
            document.add_paragraph(i,style=None)

        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(10)
        document.save(file_name)   
        print("make file " + file_name + " success..........")

    def make_raw(self, page_header_name, file_name, data_list):
        random.shuffle(data_list)
        

        with open(file_name, "a+") as f:
            for i in data_list:
                f.write(i)
        print("make file " + file_name + " success..........")

    def make_raw_obj(self):
        data_list = self.get_str()
        f = open('somefile', 'wb')
        pickle.dump(data_list, f)
        t_name = "t.docx"
        self.make_docx("headerheaderheaderheader",t_name,data_list)
        

    def task_main_userawobj(self):
        with open("template.docx", "rb") as f:
            
            data = f.read()
            f = open('somefile', 'rb')
            data_list = pickle.load(f)
            data_list = data_list[0:70]
            # 请求一次随机打乱
            for row in self.file_date_list:
                page_header_name = row[0] + " " + row[2]
                # 牧区特大桥bim信息化管理系统(客户端) V1.0 2019-04-03 源码
                file_name_1 = row[0] + " " + row[2] + " " + row[3] + " 说明书.docx"
                file_name_2 = row[0] + " " + row[2] + " 源码.docx"
                file_name_3 = row[0] + " " + row[2] + " " + row[3] + " 主要功能和技术特点.docx"

                # with open(file_name_1, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_2, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_3, "ab+") as f2:
                #     f2.write(data)
                
                self.make_docx(page_header_name,file_name_2,data_list)

                os.makedirs(page_header_name)
                
            print("finish make docx")

    def task_main(self):
        with open("template.docx", "rb") as f:
            
            data = f.read()
            data_list = self.get_str()
            # 请求一次随机打乱

            count = 0
            t_name = "t.docx"
            self.make_docx("headerheaderheaderheader",str(count)+t_name,data_list)

            # while(len(data_list) < 12 or os.path.getsize(str(count)+t_name) > 41415 or os.path.getsize(str(count)+t_name) < 351415):
            #     print("***************************************warning data list too short , so request new data" + str(count))

            #     print("***************************************warning file too big")
            #     data_list = self.get_str()
            #     count = count +1
            #     self.make_docx("headerheaderheaderheader",str(count)+t_name,data_list)
            # print("data list length is " + str(len(data_list)))

            

            for row in self.file_date_list:
                page_header_name = row[0] + " " + row[2]
                file_name_1 = row[0] + " " + row[2] + " " + row[3] + " 说明书.docx"
                file_name_2 = row[0] + " " + row[2] + " " + row[3] + " 源码.docx"
                file_name_3 = row[0] + " " + row[2] + " " + row[3] + " 主要功能和技术特点.docx"

                # with open(file_name_1, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_2, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_3, "ab+") as f2:
                #     f2.write(data)
                
                self.make_docx(page_header_name,file_name_2,data_list)
                
            print("finish make docx")

    def task_raw(self):
        with open("template.docx", "rb") as f:
            count = 0
            t_name = "t.docx"
            data = f.read()
            data_list = self.get_str()
            # 请求一次随机打乱
            self.make_docx("headerheaderheaderheader",str(count)+t_name,data_list)

            while(len(data_list) < 12 or os.path.getsize(str(count)+t_name) > 41415 or os.path.getsize(str(count)+t_name) < 351415):
                print("***************************************warning data list too short , so request new data" + str(count))

                print("***************************************warning file too big")
                data_list = self.get_str()
                count = count +1
                self.make_docx("headerheaderheaderheader",str(count)+t_name,data_list)
            print("data list length is " + str(len(data_list)))

            

            for row in self.file_date_list:
                page_header_name = row[0] + " " + row[2]
                file_name_1 = row[0] + " " + row[2] + " " + row[3] + " 说明书.docx"
                file_name_2 = row[0] + " " + row[2] + " " + row[3] + " 源码.docx"
                file_name_3 = row[0] + " " + row[2] + " " + row[3] + " 主要功能和技术特点.docx"

                # with open(file_name_1, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_2, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_3, "ab+") as f2:
                #     f2.write(data)
                
                self.make_raw(page_header_name,file_name_2,data_list)

                # os.makedirs(page_header_name)
                
            print("finish make txt")

    
    def __del__(self):
        # 清尾工作

        try:
            self.f3.close()
            print("file object close succes")
        except:
            print("file object close faield")

    def make_docx_template(self):
        with open("template.docx", "rb") as f:
            data = f.read()
            
            for row in self.file_date_list:
                page_header_name = row[0] + " " + row[1]
                file_name_1 = row[0] + " " + row[1] + " " + row[2] + " 说明书.docx"
                # file_name_2 = row[0] + " " + row[1] + " " + row[2] + " 源码.docx"
                file_name_3 = row[0] + " " + row[1] + " " + row[2] + " 主要功能和技术特点.docx"

                # with open(file_name_1, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_2, "ab+") as f2:
                #     f2.write(data)
                # with open(file_name_3, "ab+") as f2:
                #     f2.write(data)
                data_list = []
                self.make_docx(page_header_name,file_name_1,data_list)
                self.make_docx(page_header_name,file_name_3,data_list)
                
            print("finish make docx")


if __name__ == '__main__':
    NEED  = 10
    FILTER_KEYWORD = ".java"
    SEARCH_KEYWORD = "linux"
    LAN = "java"
    page_num = 35
    page_line_count = 44
    need_line_count = page_num*page_line_count

    if sys.argv[1] == "get":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        get_doc_instance.task_main()

    if sys.argv[1] == "get_raw":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        get_doc_instance.task_raw()

    if sys.argv[1] == "test_get":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        get_doc_instance.get_str()

    if sys.argv[1] == "make_template":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        get_doc_instance.make_docx_template()

    if sys.argv[1] == "test_make":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        data = {
            "123123123123",
            "12312312312319999999999999999999"
        }
        get_doc_instance.make_docx("this header", "this is file name.docx",data)

    if sys.argv[1] == "test_read":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        get_doc_instance.read_csv()

    if sys.argv[1] == "make_raw_obj":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        get_doc_instance.make_raw_obj()

    if sys.argv[1] == "raw_task":
        get_doc_instance = get_doc(NEED,FILTER_KEYWORD,SEARCH_KEYWORD,LAN)
        get_doc_instance.task_main_userawobj()
        

    print("Time used:",time.time() - start)